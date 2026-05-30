"""
Render the academic project report as a professional .docx file.

Features:
  - Cover page (title, subtitle, generated-by line, date)
  - Auto-updating Word Table of Contents field
  - Heading 1 styled section headers
  - Justified body paragraphs in Calibri 11pt
  - Inline **bold** parsing
  - Markdown-style bullet list parsing (-, *)
  - Page-number footer
"""

from __future__ import annotations

import io
import re
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


BODY_FONT = "Calibri"
HEADING_COLOR = RGBColor(0x1F, 0x35, 0x64)  # dark navy
ACCENT_COLOR = RGBColor(0x2E, 0x75, 0xB6)   # mid blue


def _set_run_font(run, name=BODY_FONT, size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    # ensure font applies to East Asian / complex scripts too
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)


def _add_page_break(document):
    p = document.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def _add_toc(document):
    """Inject a Word TOC field. Word will populate it on open (or when user hits F9)."""
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")

    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = r'TOC \o "1-2" \h \z \u'

    fldChar_sep = OxmlElement("w:fldChar")
    fldChar_sep.set(qn("w:fldCharType"), "separate")

    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click and choose 'Update Field' to populate the Table of Contents."

    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")

    r_element = run._r
    r_element.append(fldChar_begin)
    r_element.append(instrText)
    r_element.append(fldChar_sep)
    r_element.append(placeholder)
    r_element.append(fldChar_end)

    # Ask Word to refresh fields when the document opens
    settings = document.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        update_fields.set(qn("w:val"), "true")
        settings.append(update_fields)


def _add_page_numbers_footer(document, project_name: str):
    """Footer: 'Project Name | Page X of Y' centered."""
    section = document.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.text = ""

    run = paragraph.add_run(f"{project_name}  |  Page ")
    _set_run_font(run, size=9, color=RGBColor(0x80, 0x80, 0x80))

    # PAGE field
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar"); fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar"); fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin); run._r.append(instr); run._r.append(fld_end)
    _set_run_font(run, size=9, color=RGBColor(0x80, 0x80, 0x80))

    run = paragraph.add_run(" of ")
    _set_run_font(run, size=9, color=RGBColor(0x80, 0x80, 0x80))

    # NUMPAGES field
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar"); fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "NUMPAGES"
    fld_end = OxmlElement("w:fldChar"); fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin); run._r.append(instr); run._r.append(fld_end)
    _set_run_font(run, size=9, color=RGBColor(0x80, 0x80, 0x80))


def _add_bold_runs(paragraph, text: str, size=11):
    """Split text on **...** and emit bold/normal runs accordingly."""
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            _set_run_font(run, size=size, bold=True)
        else:
            run = paragraph.add_run(part)
            _set_run_font(run, size=size)


def _add_body_block(document, body_markdown: str):
    """Render a markdown-ish body block: paragraphs, bullets, bold."""
    lines = body_markdown.split("\n")
    for raw_line in lines:
        line = raw_line.rstrip()
        if not line.strip():
            continue

        bullet_match = re.match(r"^\s*[-*]\s+(.*)$", line)
        numbered_match = re.match(r"^\s*\d+[.)]\s+(.*)$", line)

        if bullet_match:
            para = document.add_paragraph(style="List Bullet")
            _add_bold_runs(para, bullet_match.group(1).strip(), size=11)
        elif numbered_match:
            para = document.add_paragraph(style="List Number")
            _add_bold_runs(para, numbered_match.group(1).strip(), size=11)
        else:
            para = document.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.space_after = Pt(8)
            _add_bold_runs(para, line.strip(), size=11)


def _build_cover_page(document, project_name: str):
    # Vertical breathing room
    for _ in range(6):
        document.add_paragraph()

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(project_name)
    _set_run_font(run, name=BODY_FONT, size=32, bold=True, color=HEADING_COLOR)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Academic Project Report")
    _set_run_font(run, name=BODY_FONT, size=18, color=ACCENT_COLOR)

    for _ in range(8):
        document.add_paragraph()

    badge = document.add_paragraph()
    badge.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = badge.add_run("Generated by Vivora")
    _set_run_font(run, name=BODY_FONT, size=12, italic=True, color=RGBColor(0x60, 0x60, 0x60))

    date_p = document.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run(datetime.now().strftime("%B %d, %Y"))
    _set_run_font(run, name=BODY_FONT, size=11, color=RGBColor(0x80, 0x80, 0x80))

    _add_page_break(document)


def _build_toc_page(document):
    heading = document.add_paragraph()
    run = heading.add_run("Table of Contents")
    _set_run_font(run, name=BODY_FONT, size=22, bold=True, color=HEADING_COLOR)
    document.add_paragraph()  # spacer
    _add_toc(document)
    _add_page_break(document)


def report_to_docx(report: dict) -> bytes:
    """
    `report` is the dict returned by generate_full_report():
        { "project_name": str, "sections": dict[str,str], "markdown": str }
    """
    project_name = report.get("project_name", "Project Report")
    sections = report.get("sections", {})

    document = Document()

    # Page margins
    for section in document.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Default body style
    normal_style = document.styles["Normal"]
    normal_style.font.name = BODY_FONT
    normal_style.font.size = Pt(11)

    # Heading 1 styling
    h1 = document.styles["Heading 1"]
    h1.font.name = BODY_FONT
    h1.font.size = Pt(18)
    h1.font.bold = True
    h1.font.color.rgb = HEADING_COLOR

    _build_cover_page(document, project_name)
    _build_toc_page(document)
    _add_page_numbers_footer(document, project_name)

    for section_name, body in sections.items():
        heading = document.add_heading(section_name, level=1)
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(8)
        _add_body_block(document, body)
        document.add_paragraph()  # spacer between sections

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
